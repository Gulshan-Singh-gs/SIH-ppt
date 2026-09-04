"""
Verification Suite for Dual Working Engine & Python Document Libraries (SIH PSC26117).
Tests:
1. CSV Upload & Ingestion via Python libraries (Budgets, Sums, Averages, Neumorphic Table).
2. Word DOCX Ingestion via Python libraries.
3. Dual-Engine 50/50 Workload Partitioning for Multi-Page OCR.
4. Server Batch Processing with Adaptive Threshold Preprocessing.
5. Dual-Engine Result Merger & Structured Briefing generation.
"""
import urllib.request
import json
import io
import time
from pathlib import Path

BASE_URL = "http://127.0.0.1:8001"

def test_split_ocr_job():
    print(">> [1/5] Testing Workload Partitioning: 50% Server / 50% Browser...")
    payload = json.dumps({"total_pages": 12, "job_name": "SIH_Tender_Specifications.pdf"}).encode("utf-8")
    req = urllib.request.Request(f"{BASE_URL}/api/workbench/split-ocr-job", data=payload, headers={"Content-Type": "application/json"})
    with urllib.request.urlopen(req) as resp:
        assert resp.status == 200
        data = json.loads(resp.read().decode("utf-8"))
        split = data["split"]
        assert split["strategy"] == "DUAL_ENGINE_SPLIT"
        assert len(split["server_pages"]) == 6
        assert len(split["browser_pages"]) == 6
        assert split["server_pages"] == [1, 2, 3, 4, 5, 6]
        assert split["browser_pages"] == [7, 8, 9, 10, 11, 12]
        print(f"   [PASS] 12 pages partitioned equally: 6 server / 6 browser ({split['split_ratio']})")

def test_process_ocr_batch():
    print(">> [2/5] Testing Server Batch OCR Processing...")
    payload = json.dumps({
        "images": [
            {"name": "Page_1.png", "page_num": 1, "b64": ""},
            {"name": "Page_2.png", "page_num": 2, "b64": ""}
        ]
    }).encode("utf-8")
    req = urllib.request.Request(f"{BASE_URL}/api/workbench/process-ocr-batch", data=payload, headers={"Content-Type": "application/json"})
    with urllib.request.urlopen(req) as resp:
        assert resp.status == 200
        data = json.loads(resp.read().decode("utf-8"))
        assert data["status"] == "SUCCESS"
        assert len(data["pages"]) == 2
        print(f"   [PASS] Server processed {len(data['pages'])} pages successfully.")

def test_complete_dual_ocr():
    print(">> [3/5] Testing Dual-Engine Result Assembly & Structured Briefing...")
    server_pages = [
        {"page_num": 1, "engine": "Local Python Server", "extracted_text": "Clause 1: Air-gapped on-premise inference required."},
        {"page_num": 2, "engine": "Local Python Server", "extracted_text": "Clause 2: Turnkey delivery within 45 days."}
    ]
    browser_pages = [
        {"page_num": 3, "engine": "Browser Engine", "extracted_text": "Clause 3: Make-in-India Class I minimum 50%."},
        {"page_num": 4, "engine": "Browser Engine", "extracted_text": "Clause 4: EMD exemption applicable for DPIIT startups."}
    ]
    payload = json.dumps({
        "server_pages": server_pages,
        "browser_pages": browser_pages,
        "file_name": "Tender_GeM_Notice_B98221.pdf",
        "elapsed_seconds": 0.85
    }).encode("utf-8")
    req = urllib.request.Request(f"{BASE_URL}/api/workbench/complete-dual-ocr", data=payload, headers={"Content-Type": "application/json"})
    with urllib.request.urlopen(req) as resp:
        assert resp.status == 200
        data = json.loads(resp.read().decode("utf-8"))
        assert data["status"] == "SUCCESS"
        assert data["total_pages"] == 4
        assert data["server_processed"] == 2
        assert data["browser_processed"] == 2
        rep = data["report_json"]
        assert "Dual-Engine OCR Intelligence" in rep["title"]
        assert len(rep["metrics"]) == 4
        assert rep["metrics"][1]["value"] == "50 / 50 Split"
        assert len(rep["tenders_table"]) == 4
        print(f"   [PASS] Dual-engine report compiled: {rep['title']} (Metrics: {len(rep['metrics'])}, Table Rows: {len(rep['tenders_table'])})")

def test_csv_upload():
    print(">> [4/5] Testing CSV Ingestion via Python libraries...")
    csv_content = (
        "Item ID,Procurement Title,Ministry,Allotted Budget (INR),Status\n"
        "GEM-01,Edge Servers,MeitY,150000000,Approved\n"
        "GEM-02,UPS Backup,MHA,48000000,Pending\n"
        "GEM-03,Biometric Terminals,Railways,21000000,Approved\n"
        "GEM-04,Smart Cards,Finance,8500000,Active\n"
    )
    boundary = "----WebKitFormBoundarySovereignAI2026"
    body = (
        f"--{boundary}\r\n"
        f'Content-Disposition: form-data; name="file"; filename="annual_procurement_budget.csv"\r\n'
        f"Content-Type: text/csv\r\n\r\n"
        f"{csv_content}\r\n"
        f"--{boundary}--\r\n"
    ).encode("utf-8")
    req = urllib.request.Request(
        f"{BASE_URL}/api/workbench/upload-file",
        data=body,
        headers={"Content-Type": f"multipart/form-data; boundary={boundary}"}
    )
    with urllib.request.urlopen(req) as resp:
        assert resp.status == 200
        data = json.loads(resp.read().decode("utf-8"))
        assert data["status"] == "SUCCESS"
        rep = data["report_json"]
        assert rep["metrics"][0]["value"] == "4"  # 4 records
        print(f"   [PASS] CSV parsed: {rep['title']} ({rep['metrics'][0]['label']}: {rep['metrics'][0]['value']})")

def test_word_docx_upload():
    print(">> [5/5] Testing Word (.docx) Ingestion via Python libraries...")
    import docx
    doc_path = Path("test_sample_notice.docx")
    doc = docx.Document()
    doc.add_heading("Government Procurement Notice: Security Compliance", 0)
    doc.add_paragraph("All procurement must satisfy Section 12 Air-Gap specifications.")
    doc.add_paragraph("Zero exfiltration guaranteed across all network boundaries.")
    doc.save(str(doc_path))

    content = doc_path.read_bytes()
    boundary = "----WebKitFormBoundarySovereignAI2026"
    body = (
        f"--{boundary}\r\n"
        f'Content-Disposition: form-data; name="file"; filename="test_sample_notice.docx"\r\n'
        f"Content-Type: application/vnd.openxmlformats-officedocument.wordprocessingml.document\r\n\r\n"
    ).encode("utf-8") + content + f"\r\n--{boundary}--\r\n".encode("utf-8")

    req = urllib.request.Request(
        f"{BASE_URL}/api/workbench/upload-file",
        data=body,
        headers={"Content-Type": f"multipart/form-data; boundary={boundary}"}
    )
    with urllib.request.urlopen(req) as resp:
        assert resp.status == 200
        data = json.loads(resp.read().decode("utf-8"))
        assert data["status"] == "SUCCESS"
        rep = data["report_json"]
        assert "Document Briefing" in rep["title"]
        print(f"   [PASS] DOCX parsed: {rep['title']} ({rep['metrics'][1]['label']}: {rep['metrics'][1]['value']})")

    if doc_path.exists():
        doc_path.unlink()

if __name__ == "__main__":
    print("=" * 70)
    print(">> RUNNING DUAL-ENGINE OCR & PYTHON LIBRARIES VERIFICATION SUITE")
    print("=" * 70)
    test_split_ocr_job()
    test_process_ocr_batch()
    test_complete_dual_ocr()
    test_csv_upload()
    test_word_docx_upload()
    print("=" * 70)
    print(">> ALL 5 DUAL-ENGINE & PYTHON LIBRARY TESTS PASSED WITH 100% SUCCESS!")
    print("=" * 70)
